import os
from docx import Document as DocxDocument
import PyPDF2


def parse_docx(file_path):
    """
    Parse DOCX file, extract text and table content
    :param file_path: DOCX file path
    :return: List containing text and table chunks
    """
    doc = DocxDocument(file_path)
    chunks = []

    # Extract paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            chunks.append({"type": "text", "content": text})

    # Extract tables
    for table in doc.tables:
        md_table = []
        if table.rows:
            header = [cell.text.strip() for cell in table.rows[0].cells]
            md_table.append("| " + " | ".join(header) + " |")
            md_table.append("| " + " | ".join(["---"] * len(header)) + " |")
            for row in table.rows[1:]:
                row_data = [cell.text.strip() for cell in row.cells]
                md_table.append("| " + " | ".join(row_data) + " |")
            table_content = "\n".join(md_table)
            if table_content.strip():
                chunks.append({"type": "table", "content": table_content})
    return chunks


def parse_pdf(file_path):
    """
    Parse PDF file and extract text content
    :param file_path: PDF file path
    :return: List containing text chunks
    """
    chunks = []
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text = page.extract_text()
                if text and text.strip():
                    paragraphs = text.split('\n\n')
                    for para in paragraphs:
                        para = para.strip()
                        if para:
                            chunks.append({"type": "text", "content": para})
    except Exception as e:
        print(f"PDF parsing error {file_path}: {e}")
    return chunks


def parse_txt(file_path):
    """
    Parse TXT file and extract text content
    :param file_path: TXT file path
    :return: List containing text chunks
    """
    chunks = []
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
            if content.strip():
                paragraphs = content.split('\n\n')
                for para in paragraphs:
                    para = para.strip()
                    if para:
                        chunks.append({"type": "text", "content": para})
    except Exception as e:
        print(f"TXT parsing error {file_path}: {e}")
        # Try other encoding
        try:
            with open(file_path, 'r', encoding='gbk') as file:
                content = file.read()
                if content.strip():
                    paragraphs = content.split('\n\n')
                    for para in paragraphs:
                        para = para.strip()
                        if para:
                            chunks.append({"type": "text", "content": para})
        except Exception as e2:
            print(f"TXT parsing error (GBK) {file_path}: {e2}")
    return chunks


def parse_document(file_path):
    """
    Parse document based on file extension
    :param file_path: Document file path
    :return: List containing text and table chunks
    """
    ext = os.path.splitext(file_path)[1].lower()
    if ext == '.docx':
        return parse_docx(file_path)
    elif ext == '.pdf':
        return parse_pdf(file_path)
    elif ext == '.txt':
        return parse_txt(file_path)
    else:
        print(f"Unsupported file format: {ext}")
        return []


def get_all_files_in_directory(directory, extensions=None):
    """
    Recursively get all files in directory and its subdirectories
    :param directory: Directory path
    :param extensions: List of file extensions to find, e.g., ['.docx', '.pdf'], if None returns all files
    :return: List of file paths
    """
    if extensions is None:
        extensions = []

    normalized_extensions = []
    for ext in extensions:
        if not ext.startswith('.'):
            ext = '.' + ext
        normalized_extensions.append(ext.lower())

    files = []
    for root, dirs, filenames in os.walk(directory):
        for filename in filenames:
            if filename.startswith('.'):
                continue  # Skip hidden files
            if extensions:
                _, ext = os.path.splitext(filename)
                if ext.lower() in normalized_extensions:
                    files.append(os.path.join(root, filename))
            else:
                files.append(os.path.join(root, filename))
    return files